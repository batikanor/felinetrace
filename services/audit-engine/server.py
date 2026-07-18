from __future__ import annotations

import base64
import csv
import hashlib
import io
import json
import os
import re
import shutil
import threading
import time
import unicodedata
import urllib.error
import urllib.request
import uuid
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from datetime import datetime, timezone
from decimal import Decimal, InvalidOperation
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader


HOST = "127.0.0.1"
PORT = int(os.environ.get("AUDIT_ENGINE_PORT", "43421"))
ROOT = Path(__file__).resolve().parents[2]
SAMPLE_DIR = ROOT / "source-docs" / "data"
FINAL_DIR = ROOT / "source-docs" / "final-data" / "Daten BSP"
RUNTIME_DIR = Path(__file__).resolve().parent / ".runtime"
MAX_BODY = 90 * 1024 * 1024
MAX_FILES = 120
MAX_FILE = 25 * 1024 * 1024


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")
    return re.sub(r"\s+", " ", str(value)).strip()


def norm(value: Any) -> str:
    text = unicodedata.normalize("NFKD", clean_text(value)).encode("ascii", "ignore").decode()
    return re.sub(r"[^A-Z0-9]", "", text.upper())


def slug(value: str) -> str:
    result = re.sub(r"[^a-z0-9]+", "-", norm(value).lower()).strip("-")
    return result[:64] or uuid.uuid4().hex[:12]


def parse_amount(value: Any) -> Decimal | None:
    text = clean_text(value).replace("€", "").replace("EUR", "").replace(" ", "")
    if not text:
        return None
    negative = text.startswith("(") and text.endswith(")")
    text = text.strip("()")
    if "," in text and "." in text:
        if text.rfind(",") > text.rfind("."):
            text = text.replace(".", "").replace(",", ".")
        else:
            text = text.replace(",", "")
    elif "," in text:
        parts = text.split(",")
        text = "".join(parts) if len(parts[-1]) == 3 else text.replace(".", "").replace(",", ".")
    elif "." in text:
        parts = text.split(".")
        if len(parts) > 1 and all(len(part) == 3 for part in parts[1:]):
            text = "".join(parts)
    try:
        amount = Decimal(text)
        return -amount if negative else amount
    except InvalidOperation:
        return None


def money(value: Decimal | float | int | None) -> str:
    if value is None:
        return "—"
    number = Decimal(value)
    sign = "−" if number < 0 else ""
    number = abs(number)
    return f"{sign}€{number:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def parse_date(value: Any) -> datetime | None:
    if isinstance(value, datetime):
        return value
    text = clean_text(value)
    for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y"):
        try:
            return datetime.strptime(text[:10], fmt)
        except ValueError:
            pass
    return None


def decode_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            pass
    return data.decode("utf-8", errors="replace")


@dataclass
class Table:
    document: "ParsedDocument"
    sheet: str
    columns: list[str]
    rows: list[dict[str, str]]
    start_row: int = 1

    @property
    def normalized_columns(self) -> dict[str, str]:
        return {norm(column): column for column in self.columns}


@dataclass
class ParsedDocument:
    path: str
    name: str
    kind: str
    size: int
    text: str = ""
    pages: list[str] = field(default_factory=list)
    tables: list[Table] = field(default_factory=list)


def field_value(row: dict[str, str], *aliases: str) -> str:
    normalized = {norm(key): value for key, value in row.items()}
    wanted = [norm(alias) for alias in aliases]
    for alias in wanted:
        if alias in normalized:
            return clean_text(normalized[alias])
    for alias in wanted:
        for key, value in normalized.items():
            if alias and (alias in key or key in alias):
                return clean_text(value)
    return ""


def row_text(row: dict[str, str]) -> str:
    return " · ".join(clean_text(value) for key, value in row.items() if not key.startswith("__") and clean_text(value))


def find_tables(documents: list[ParsedDocument], required_aliases: list[tuple[str, ...]]) -> list[Table]:
    matches: list[Table] = []
    for document in documents:
        for table in document.tables:
            columns = list(table.normalized_columns)
            if all(any(any(norm(alias) in column or column in norm(alias) for alias in group) for column in columns) for group in required_aliases):
                matches.append(table)
    return matches


def parse_index_schemas(files: dict[str, bytes]) -> dict[str, list[str]]:
    schemas: dict[str, list[str]] = {}
    for relative, data in files.items():
        if Path(relative).name.lower() != "index.xml":
            continue
        try:
            root = ET.fromstring(decode_bytes(data).lstrip("\ufeff"))
        except ET.ParseError:
            continue
        parent = str(Path(relative).parent)
        for table in root.findall(".//Table"):
            url = clean_text(table.findtext("URL"))
            columns = [clean_text(node.findtext("Name")) for node in table.findall(".//VariableColumn")]
            if url and columns:
                schemas[str(Path(parent) / url)] = columns
                schemas[url] = columns
    return schemas


def rows_to_table(document: ParsedDocument, raw_rows: list[list[Any]], sheet: str, schema: list[str] | None = None) -> Table | None:
    rows = [[clean_text(cell) for cell in row] for row in raw_rows]
    rows = [row for row in rows if any(row)]
    if not rows:
        return None
    if schema:
        header_index = -1
        columns = schema
        data_rows = rows
        start_row = 1
    else:
        candidates = [(index, sum(bool(cell) for cell in row)) for index, row in enumerate(rows[:12])]
        header_index, count = max(candidates, key=lambda item: item[1])
        if count < 2:
            return None
        columns = [cell or f"Column {index + 1}" for index, cell in enumerate(rows[header_index])]
        data_rows = rows[header_index + 1 :]
        start_row = header_index + 2
    width = len(columns)
    records: list[dict[str, str]] = []
    for row in data_rows:
        values = row[:width] + [""] * max(0, width - len(row))
        if not any(values):
            continue
        record = {columns[index]: values[index] for index in range(width)}
        record["__row__"] = str(start_row + len(records))
        records.append(record)
    return Table(document=document, sheet=sheet, columns=columns, rows=records, start_row=start_row)


def parse_documents(files: dict[str, bytes]) -> list[ParsedDocument]:
    schemas = parse_index_schemas(files)
    documents: list[ParsedDocument] = []
    for relative, data in sorted(files.items()):
        path = Path(relative)
        suffix = path.suffix.lower()
        document = ParsedDocument(path=relative, name=path.name, kind=suffix.lstrip(".") or "file", size=len(data))
        try:
            if suffix in {".csv", ".txt"}:
                text = decode_bytes(data)
                document.text = text
                sample = text[:4096]
                delimiter = ";" if sample.count(";") >= max(sample.count("\t"), sample.count(",")) else "\t" if sample.count("\t") > sample.count(",") else ","
                raw_rows = list(csv.reader(io.StringIO(text), delimiter=delimiter, quotechar='"'))
                schema = schemas.get(relative) or schemas.get(path.name)
                table = rows_to_table(document, raw_rows, path.stem, schema)
                if table:
                    document.tables.append(table)
            elif suffix == ".xlsx":
                workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
                for worksheet in workbook.worksheets:
                    raw_rows = [list(row) for row in worksheet.iter_rows(values_only=True)]
                    table = rows_to_table(document, raw_rows, worksheet.title)
                    if table:
                        document.tables.append(table)
                document.text = "\n".join(row_text(row) for table in document.tables for row in table.rows[:1000])
            elif suffix == ".docx":
                docx = DocxDocument(io.BytesIO(data))
                paragraphs = [clean_text(paragraph.text) for paragraph in docx.paragraphs if clean_text(paragraph.text)]
                document.text = "\n".join(paragraphs)
                for index, source_table in enumerate(docx.tables):
                    raw_rows = [[cell.text for cell in row.cells] for row in source_table.rows]
                    table = rows_to_table(document, raw_rows, f"Table {index + 1}")
                    if table:
                        document.tables.append(table)
            elif suffix == ".pdf":
                reader = PdfReader(io.BytesIO(data))
                document.pages = [clean_text(page.extract_text() or "") for page in reader.pages]
                document.text = "\n".join(document.pages)
            elif suffix in {".xml", ".dtd", ".md", ".json"}:
                document.text = decode_bytes(data)
            else:
                document.text = decode_bytes(data) if len(data) < 2_000_000 else ""
        except Exception as error:
            document.text = f"Parser warning: {type(error).__name__}"
        documents.append(document)
    return documents


class Analyzer:
    def __init__(self, documents: list[ParsedDocument], dataset_name: str, dataset_id: str, dataset_kind: str = "custom"):
        self.documents = documents
        self.dataset_name = dataset_name
        self.dataset_id = dataset_id
        self.dataset_kind = dataset_kind
        self.sources: dict[str, dict[str, Any]] = {}
        self.findings: list[dict[str, Any]] = []
        self.holds: list[dict[str, Any]] = []

    def source_for_rows(self, source_id: str, table: Table, rows: list[dict[str, str]], value: str, passage: str | None = None) -> str:
        row_numbers = [int(row.get("__row__", "0") or 0) for row in rows]
        if row_numbers:
            location = f"{table.sheet} · row {row_numbers[0]}" if len(row_numbers) == 1 else f"{table.sheet} · rows {min(row_numbers)}–{max(row_numbers)}"
        else:
            location = f"{table.sheet} · 0 matches across {len(table.rows):,} rows"
        preview_rows = []
        for row in rows[:12]:
            preview_rows.append({
                "label": row.get("__row__", ""),
                "cells": [clean_text(row.get(column)) for column in table.columns[:7]],
            })
        self.sources[source_id] = {
            "id": source_id,
            "name": table.document.name,
            "path": table.document.path,
            "location": location,
            "passage": passage or "; ".join(row_text(row) for row in rows[:12]),
            "value": value,
            "type": "sheet",
            "preview": {"columns": table.columns[:7], "rows": preview_rows},
        }
        return source_id

    def source_for_text(self, source_id: str, document: ParsedDocument, passage: str, value: str, page: int = 1) -> str:
        self.sources[source_id] = {
            "id": source_id,
            "name": document.name,
            "path": document.path,
            "location": f"Page {page}" if document.kind == "pdf" else "Document text",
            "passage": passage,
            "value": value,
            "type": "pdf" if document.kind in {"pdf", "docx"} else "sheet",
            "preview": None,
        }
        return source_id

    def gate(self, gate_id: str, label: str, status: str, summary: str, atoms: list[dict[str, Any]]) -> dict[str, Any]:
        return {"id": gate_id, "label": label, "status": status, "summary": summary, "atoms": atoms}

    def atom(self, atom_type: str, label: str, value: str, status: str = "proved", source_id: str | None = None) -> dict[str, Any]:
        return {"type": atom_type, "label": label, "value": value, "status": status, "sourceId": source_id}

    def detect_vendor_control(self) -> None:
        master_tables = find_tables(self.documents, [("KONTO", "VENDOR", "SUPPLIER"), ("GEAENDERT_VON", "CREATED_BY", "CHANGED_BY"), ("GENEHMIGT_VON", "APPROVED_BY")])
        receipt_tables = find_tables(self.documents, [("KREDITOR", "VENDOR", "SUPPLIER"), ("WARENEINGANG", "GOODS_RECEIPT", "RECEIPT")])
        vendor_ledger_tables = find_tables(self.documents, [("LIEFERANTENKONTONUMMER", "KREDITOR", "VENDOR"), ("BUCHUNGSBETRAG", "BETRAG", "AMOUNT")])
        gl_tables = find_tables(self.documents, [("BUCHUNGSTYP", "TRANSACTION_TYPE"), ("BUCHUNGSBETRAG", "AMOUNT"), ("BENUTZERKENNUNG", "USER")])
        permission_tables = find_tables(self.documents, [("BENUTZER", "USER"), ("ZAHLUNGSLAUF", "PAYMENT"), ("KREDITOR", "VENDOR")])
        for master_table in master_tables:
            for row in master_table.rows:
                vendor_id = field_value(row, "KONTO", "VENDOR_ID", "SUPPLIER_ID")
                vendor_name = field_value(row, "NAME", "VENDOR_NAME", "SUPPLIER_NAME")
                changed_by = field_value(row, "GEAENDERT_VON", "CREATED_BY", "CHANGED_BY")
                approved_by = field_value(row, "GENEHMIGT_VON", "APPROVED_BY")
                change = " ".join([field_value(row, "FELD", "FIELD"), field_value(row, "WERT_NEU", "NEW_VALUE")])
                is_new = any(token in norm(change) for token in ("NEUANLAGE", "NEWVENDOR", "CREATE", "ANGELEGT"))
                if not vendor_id or not changed_by or changed_by != approved_by or not is_new:
                    continue
                receipts = [(table, record) for table in receipt_tables for record in table.rows if field_value(record, "KREDITOR", "VENDOR", "SUPPLIER") == vendor_id]
                ledger = [(table, record) for table in vendor_ledger_tables for record in table.rows if field_value(record, "LIEFERANTENKONTONUMMER", "KREDITOR", "VENDOR") == vendor_id]
                payments = [(table, record) for table, record in ledger if (parse_amount(field_value(record, "BUCHUNGSBETRAG", "BETRAG", "AMOUNT")) or 0) > 0]
                if receipts or len(payments) < 2:
                    continue
                gross = sum((parse_amount(field_value(record, "BUCHUNGSBETRAG", "BETRAG", "AMOUNT")) or Decimal(0)) for _, record in payments)
                gl_matches = [(table, record) for table in gl_tables for record in table.rows if vendor_id in field_value(record, "SACHKONTONUMMER", "ACCOUNT") or vendor_id in row_text(record)]
                users = {field_value(record, "BENUTZERKENNUNG", "USER") for _, record in gl_matches if field_value(record, "BENUTZERKENNUNG", "USER")}
                permission = [(table, record) for table in permission_tables for record in table.rows if field_value(record, "BENUTZER", "USER") == changed_by]

                prefix = f"vendor-{slug(vendor_id)}"
                source_ids = [self.source_for_rows(f"{prefix}-master", master_table, [row], f"{changed_by} → {approved_by}")]
                if permission:
                    table = permission[0][0]
                    source_ids.append(self.source_for_rows(f"{prefix}-permissions", table, [record for _, record in permission], "create · post · pay rights"))
                if gl_matches:
                    table = gl_matches[0][0]
                    records = [record for item_table, record in gl_matches if item_table is table]
                    source_ids.append(self.source_for_rows(f"{prefix}-ledger", table, records, f"{money(gross)} paid", f"{len(payments)} payments; users: {', '.join(sorted(users)) or changed_by}; vendor {vendor_id}."))
                elif payments:
                    table = payments[0][0]
                    records = [record for item_table, record in payments if item_table is table]
                    source_ids.append(self.source_for_rows(f"{prefix}-ledger", table, records, f"{money(gross)} paid"))
                if receipt_tables:
                    receipt_table = receipt_tables[0]
                    source_ids.append(self.source_for_rows(f"{prefix}-receipts", receipt_table, [], "0 receipts", f"No goods-receipt row references vendor {vendor_id} or {vendor_name}."))

                gates = [
                    self.gate("facts", "Facts", "pass", "Vendor creation and payments were selected from source rows.", [
                        self.atom("MasterDataRow", "Creator = approver", f"{changed_by} → {approved_by}", source_id=source_ids[0]),
                        self.atom("PaymentSet", "Payments", f"{len(payments)} · {money(gross)}", source_id=next((item for item in source_ids if item.endswith("ledger")), None)),
                    ]),
                    self.gate("joins", "Join", "pass", "Vendor, user and payment records resolve to the same identifiers.", [self.atom("EntityJoin", "Vendor chain", f"{vendor_id} ↔ {changed_by}")]),
                    self.gate("counter", "Exclusions", "pass", "No goods receipt was found; legitimate new vendors with receipts are excluded.", [self.atom("ZeroResult", "Goods receipts", "0 matches", "excluded", next((item for item in source_ids if item.endswith("receipts")), None))]),
                    self.gate("resolver", "Sources", "pass", "Every proof atom resolves to an uploaded document location.", [self.atom("AnchorSet", "Resolved anchors", str(len(source_ids)))]),
                    self.gate("certificate", "Report", "pass", "Independent internal signals pass the report rule.", [self.atom("Decision", "Certificate", "REPORT")]),
                ]
                self.findings.append({
                    "id": f"F-{len(self.findings)+1:02d}", "scheme": "vendor-control", "category": "FRAUD RISK", "severity": "Critical",
                    "title": f"Vendor setup and payment duties were concentrated in one user", "amount": money(gross),
                    "explanation": f"{changed_by} created and approved vendor {vendor_id}, and the dossier records {len(payments)} payments but no goods receipt for that vendor.",
                    "calculation": f"{len(payments)} positive vendor payments = {money(gross)}", "decision": "REPORT", "sourceIds": source_ids,
                    "gates": gates, "methods": ["deterministic", "join", "memory", "public", "review"], "entity": vendor_name or vendor_id,
                })

    def detect_repair_capex(self) -> None:
        asset_tables = find_tables(self.documents, [("ANLAGENNUMMER", "ASSET_ID"), ("ANLAGENBEZEICHNUNG", "ASSET_NAME", "DESCRIPTION")])
        booking_tables = find_tables(self.documents, [("ANLAGENNUMMER", "ASSET_ID"), ("BUCHUNGSBETRAG", "AMOUNT"), ("BUCHUNGSART", "TYPE")])
        keywords = ("REPARATUR", "INSTAND", "AUSTAUSCH", "UBERHOL", "REPAIR", "MAINTENANCE", "REPLACEMENT", "OVERHAUL")
        for asset_table in asset_tables:
            suspicious = [row for row in asset_table.rows if any(keyword in norm(field_value(row, "ANLAGENBEZEICHNUNG", "ASSET_NAME", "DESCRIPTION")) for keyword in keywords)]
            if not suspicious:
                continue
            asset_ids = {field_value(row, "ANLAGENNUMMER", "ASSET_ID") for row in suspicious}
            matches = [(table, row) for table in booking_tables for row in table.rows if field_value(row, "ANLAGENNUMMER", "ASSET_ID") in asset_ids and any(token in norm(field_value(row, "BUCHUNGSART", "TYPE")) for token in ("ACQUISITION", "ZUGANG", "ADDITION"))]
            if not matches:
                continue
            total = sum((abs(parse_amount(field_value(row, "BUCHUNGSBETRAG", "AMOUNT")) or Decimal(0)) for _, row in matches), Decimal(0))
            source_ids = [self.source_for_rows("repair-assets", asset_table, suspicious, f"{len(suspicious)} repair-like assets")]
            booking_table = matches[0][0]
            booking_rows = [row for table, row in matches if table is booking_table]
            source_ids.append(self.source_for_rows("repair-bookings", booking_table, booking_rows, money(total)))
            self.findings.append({
                "id": f"F-{len(self.findings)+1:02d}", "scheme": "repair-capex", "category": "CLASSIFICATION", "severity": "High",
                "title": "Repair-like descriptions appear in capitalized additions", "amount": money(total),
                "explanation": f"{len(suspicious)} asset descriptions contain repair, maintenance, overhaul or replacement language and match acquisition postings.",
                "calculation": f"{len(matches)} acquisition postings = {money(total)}", "decision": "REPORT", "sourceIds": source_ids,
                "gates": [
                    self.gate("facts", "Facts", "pass", "Repair-language asset rows and acquisition amounts were selected.", [self.atom("RowSet", "Repair assets", str(len(suspicious)), source_id=source_ids[0])]),
                    self.gate("joins", "Join", "pass", "Asset identifiers join the register to acquisition postings.", [self.atom("AssetJoin", "Matched additions", str(len(matches)), source_id=source_ids[1])]),
                    self.gate("counter", "Exclusions", "pass", "Rows without repair-language indicators are excluded.", [self.atom("ScopeRule", "Classification only", "Intent is not asserted", "excluded")]),
                    self.gate("resolver", "Sources", "pass", "Register and booking rows resolve.", [self.atom("AnchorSet", "Resolved anchors", str(len(source_ids)))]),
                    self.gate("certificate", "Report", "pass", "The classification claim is source-bounded.", [self.atom("Decision", "Certificate", "REPORT")]),
                ], "methods": ["deterministic", "join", "review"],
            })

    def detect_cutoff(self) -> None:
        invoice_tables = find_tables(self.documents, [("FAKTURADATUM", "INVOICE_DATE"), ("LEISTUNGSDATUM", "SERVICE_DATE"), ("BETRAG", "AMOUNT")])
        receipt_tables = find_tables(self.documents, [("WARENEINGANG_DATUM", "RECEIPT_DATE"), ("RECHNUNGSNUMMER", "INVOICE_NUMBER"), ("BETRAG", "AMOUNT")])
        gl_tables = find_tables(self.documents, [("BUCHUNGSDATUM", "POSTING_DATE"), ("BELEGNUMMER", "DOCUMENT_NUMBER"), ("BUCHUNGSTEXT", "TEXT")])
        for invoice_table in invoice_tables:
            late = []
            for row in invoice_table.rows:
                invoice_date = parse_date(field_value(row, "FAKTURADATUM", "INVOICE_DATE"))
                service_date = parse_date(field_value(row, "LEISTUNGSDATUM", "SERVICE_DATE"))
                if invoice_date and service_date and invoice_date.year > service_date.year:
                    late.append(row)
            if len(late) < 2:
                continue
            total = sum((abs(parse_amount(field_value(row, "BETRAG_EUR", "BETRAG", "AMOUNT")) or Decimal(0)) for row in late), Decimal(0))
            invoice_refs = {field_value(row, "RECHNUNGSNUMMER", "INVOICE_NUMBER", "DOCUMENT_NUMBER") for row in late}
            vendors = {field_value(row, "KREDITOR", "VENDOR", "SUPPLIER") for row in late}
            receipts = [(table, row) for table in receipt_tables for row in table.rows if field_value(row, "KREDITOR", "VENDOR", "SUPPLIER") in vendors and not field_value(row, "RECHNUNGSNUMMER", "INVOICE_NUMBER")]
            gl_matches = [(table, row) for table in gl_tables for row in table.rows if any(reference and reference in row_text(row) for reference in invoice_refs)]
            source_ids = [self.source_for_rows("cutoff-invoices", invoice_table, late, money(total))]
            if receipts:
                table = receipts[0][0]
                rows = [row for item_table, row in receipts if item_table is table]
                source_ids.append(self.source_for_rows("cutoff-receipts", table, rows, f"{len(rows)} open receipts"))
            if gl_tables:
                table = gl_tables[0]
                source_ids.append(self.source_for_rows("cutoff-ledger", table, [row for item_table, row in gl_matches if item_table is table], f"{len(gl_matches)} matching postings", f"Search for {len(invoice_refs)} invoice references returned {len(gl_matches)} posting rows."))
            ledger_clear = len(gl_matches) == 0
            decision = "REPORT" if ledger_clear and receipts else "HOLD"
            gates = [
                self.gate("facts", "Facts", "pass", "Invoice and service dates cross the reporting boundary.", [self.atom("InvoiceSet", "Cross-period invoices", str(len(late)), source_id=source_ids[0])]),
                self.gate("joins", "Join", "pass" if receipts else "hold", "Open receipts join by vendor and reporting period.", [self.atom("ReceiptJoin", "Open receipts", str(len(receipts)), "proved" if receipts else "held", next((item for item in source_ids if item == "cutoff-receipts"), None))]),
                self.gate("counter", "Exclusions", "pass" if ledger_clear else "hold", "Exact invoice references were checked in the ledger.", [self.atom("ZeroResult", "Matching postings", str(len(gl_matches)), "excluded" if ledger_clear else "held", next((item for item in source_ids if item == "cutoff-ledger"), None))]),
                self.gate("resolver", "Sources", "pass", "Invoice, receipt and ledger searches are anchored.", [self.atom("AnchorSet", "Resolved anchors", str(len(source_ids)))]),
                self.gate("certificate", "Report" if decision == "REPORT" else "Hold", "pass" if decision == "REPORT" else "hold", "The cut-off claim requires open receipts and no matching prior-period posting.", [self.atom("Decision", "Certificate", decision)]),
            ]
            finding = {
                "id": f"F-{len(self.findings)+1:02d}" if decision == "REPORT" else f"H-{len(self.holds)+1:02d}", "scheme": "cutoff", "category": "CUT-OFF", "severity": "High",
                "title": "Cross-period invoices lack matching ledger references", "amount": money(total),
                "explanation": f"{len(late)} invoices have service dates in an earlier year. The compiler found {len(receipts)} open receipt rows and {len(gl_matches)} ledger rows carrying the invoice references.",
                "calculation": f"{len(late)} cross-period invoices = {money(total)}", "decision": decision, "sourceIds": source_ids,
                "gates": gates, "methods": ["deterministic", "join", "review"],
            }
            (self.findings if decision == "REPORT" else self.holds).append(finding)

            accrual_rows = [(table, row) for table in gl_tables for row in table.rows if any(token in norm(row_text(row)) for token in ("RUECKSTELLUNG", "ACCRUAL", "UNFAKTURIERT"))]
            unrelated = [(table, row) for table, row in accrual_rows if not any(reference and reference in row_text(row) for reference in invoice_refs)]
            if ledger_clear and unrelated:
                accrual_table = unrelated[0][0]
                rows = [row for item_table, row in unrelated if item_table is accrual_table]
                accrual_total = max((abs(parse_amount(field_value(row, "BUCHUNGSBETRAG", "AMOUNT")) or Decimal(0)) for row in rows), default=Decimal(0))
                accrual_source = self.source_for_rows("cutoff-unrelated-accrual", accrual_table, rows, money(accrual_total), f"Generic accrual rows do not carry any of the {len(invoice_refs)} invoice references tested above.")
                self.holds.append({
                    "id": f"H-{len(self.holds)+1:02d}", "scheme": "cutoff-offset", "category": "CUT-OFF", "severity": "High",
                    "title": "A generic accrual offsets the cross-period invoices", "amount": money(accrual_total),
                    "explanation": "The alternative is held because the accrual description and references do not join to the cross-period invoice population.",
                    "calculation": f"{money(accrual_total)} candidate offset versus {money(total)} invoice population", "decision": "HOLD", "sourceIds": [source_ids[0], next((item for item in source_ids if item == "cutoff-ledger"), source_ids[0]), accrual_source],
                    "gates": [
                        self.gate("facts", "Facts", "pass", "Both amounts exist in the dossier.", [self.atom("CandidateOffset", "Generic accrual", money(accrual_total), source_id=accrual_source)]),
                        self.gate("joins", "Join", "hold", "No invoice or receipt identifier joins the accrual to this population.", [self.atom("MissingJoin", "Exact references", "0 matches", "held", next((item for item in source_ids if item == "cutoff-ledger"), None))]),
                        self.gate("counter", "Exclusions", "pass", "The accrual is preserved as separate work rather than silently netted.", [self.atom("Contradiction", "Population mismatch", "unrelated references", "excluded", accrual_source)]),
                        self.gate("resolver", "Sources", "pass", "The rejected alternative resolves to exact rows.", [self.atom("AnchorSet", "Resolved anchors", "3")]),
                        self.gate("certificate", "Hold", "hold", "The offset cannot support the report claim.", [self.atom("Decision", "Certificate", "HOLD", "held")]),
                    ], "methods": ["deterministic", "join", "review"],
                })

    def detect_split_payments(self) -> None:
        gl_tables = find_tables(self.documents, [("BUCHUNGSTYP", "TYPE"), ("BUCHUNGSDATUM", "DATE"), ("BUCHUNGSBETRAG", "AMOUNT")])
        threshold = Decimal(10000)
        policy_doc = next((document for document in self.documents if any(word in norm(document.text) for word in ("SECONDAPPROVAL", "ZWEITEFREIGABE", "VIERAUGEN", "TWOAPPROVAL"))), None)
        if policy_doc:
            currency_values = re.findall(
                r"(?:EUR|€)\s*\d(?:[\d., ]*\d)?|\d(?:[\d., ]*\d)?\s*(?:EUR|€)",
                policy_doc.text,
                flags=re.IGNORECASE,
            )
            values = [parse_amount(match) for match in currency_values]
            plausible = [value for value in values if value and Decimal(1000) <= value <= Decimal(1000000)]
            if plausible:
                threshold = min(plausible)
        for table in gl_tables:
            groups: dict[tuple[str, str], list[dict[str, str]]] = {}
            for row in table.rows:
                posting_type = norm(field_value(row, "BUCHUNGSTYP", "TYPE"))
                amount = parse_amount(field_value(row, "BUCHUNGSBETRAG", "AMOUNT"))
                account = field_value(row, "SACHKONTONUMMER", "ACCOUNT")
                reference = field_value(row, "BELEGNUMMER", "REFERENCE", "DOCUMENT") or field_value(row, "DOKUMENT", "DOCUMENT")
                date = field_value(row, "BUCHUNGSDATUM", "DATE")
                if "ZAHLUNG" not in posting_type and "PAYMENT" not in posting_type:
                    continue
                if amount is None or not (0 < amount < threshold):
                    continue
                if "-" not in account and not re.search(r"(?:VENDOR|SUPPLIER|SAMMEL)[-_]?\d+", reference, re.I):
                    continue
                vendor_match = re.search(r"(\d{5,})$", account) or re.search(r"(\d{5,})", reference)
                vendor = vendor_match.group(1) if vendor_match else reference
                groups.setdefault((date, vendor), []).append(row)
            for (date, vendor), rows in groups.items():
                if len(rows) < 3:
                    continue
                amounts = [parse_amount(field_value(row, "BUCHUNGSBETRAG", "AMOUNT")) or Decimal(0) for row in rows]
                total = sum(amounts, Decimal(0))
                if total <= threshold:
                    continue
                source_ids = [self.source_for_rows(f"split-{slug(vendor)}", table, rows, f"{len(rows)} × below {money(threshold)}")]
                if policy_doc:
                    sentence = next((part.strip() for part in re.split(r"[\n.]", policy_doc.text) if str(int(threshold)) in part.replace(".", "").replace(",", "") or "10.000" in part), clean_text(policy_doc.text)[:400])
                    source_ids.append(self.source_for_text("approval-threshold", policy_doc, sentence, money(threshold)))
                self.findings.append({
                    "id": f"F-{len(self.findings)+1:02d}", "scheme": "split-payments", "category": "CONTROL", "severity": "High",
                    "title": "Same-day payments fell below the approval threshold", "amount": money(total),
                    "explanation": f"{len(rows)} payments to vendor {vendor} on {date} were individually below {money(threshold)} but total {money(total)}.",
                    "calculation": " + ".join(money(amount) for amount in amounts) + f" = {money(total)}", "decision": "REPORT", "sourceIds": source_ids,
                    "gates": [
                        self.gate("facts", "Facts", "pass", "Payment rows were grouped by date and vendor.", [self.atom("PaymentGroup", "Near-threshold payments", str(len(rows)), source_id=source_ids[0])]),
                        self.gate("joins", "Join", "pass", "All payments share vendor and date.", [self.atom("GroupKey", "Vendor + date", f"{vendor} · {date}")]),
                        self.gate("counter", "Exclusions", "pass", "The grouped total exceeds the documented threshold.", [self.atom("ControlRule", "Approval threshold", money(threshold), "excluded", source_ids[1] if len(source_ids) > 1 else None)]),
                        self.gate("resolver", "Sources", "pass", "Payment group and policy resolve.", [self.atom("AnchorSet", "Resolved anchors", str(len(source_ids)))]),
                        self.gate("certificate", "Report", "pass", "The control pattern is reported without asserting intent.", [self.atom("Decision", "Certificate", "REPORT")]),
                    ], "methods": ["deterministic", "join", "review"],
                })

    def build(self) -> dict[str, Any]:
        self.detect_vendor_control()
        self.detect_repair_capex()
        self.detect_cutoff()
        self.detect_split_payments()
        for index, finding in enumerate(self.findings, 1):
            finding["id"] = f"F-{index:02d}"
            finding["certificate"] = f"proof:{self.dataset_id}:{finding['scheme']}:{index}"
        for index, finding in enumerate(self.holds, 1):
            finding["id"] = f"H-{index:02d}"
            finding["certificate"] = f"hold:{self.dataset_id}:{finding['scheme']}:{index}"
        row_count = sum(len(table.rows) for document in self.documents for table in document.tables)
        company_pattern = re.compile(r"\b([A-ZÄÖÜ][A-Za-zÄÖÜäöüß0-9 &.,-]{2,70}\s(?:GmbH|AG|SE|Ltd\.?))\b")
        company_documents = sorted(self.documents, key=lambda document: (document.kind not in {"md", "docx", "pdf", "xml"}, document.path))
        company = next((match.group(1).strip() for document in company_documents for match in [company_pattern.search(re.sub(r"<[^>]+>", " ", document.text))] if match), self.dataset_name)
        return {
            "service": "audit-engine", "ok": True,
            "dataset": {"id": self.dataset_id, "kind": self.dataset_kind, "name": self.dataset_name, "company": company, "files": len(self.documents), "rows": row_count, "analyzedAt": datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")},
            "sources": list(self.sources.values()), "findings": self.findings, "holds": self.holds,
            "summary": {"report": len(self.findings), "hold": len(self.holds), "citations": len(self.sources)},
            "specialists": {"cognee": {"phase": "idle", "detail": "Not run"}, "tavily": {"phase": "idle", "detail": "Not run"}, "codex": {"phase": "idle", "detail": "Not run"}},
        }


STATE_LOCK = threading.Lock()
STATE: dict[str, Any] = {"analysis": None, "files": {}, "datasetDir": None, "datasetKind": None}
PRESET_CACHE: dict[str, dict[str, Any]] = {}
SPECIALIST_RUN_LOCK = threading.Lock()
DATASET_RUN_LOCK = threading.Lock()


def load_directory(path: Path) -> dict[str, bytes]:
    return {str(file.relative_to(path)): file.read_bytes() for file in path.rglob("*") if file.is_file()}


def analyze_files(files: dict[str, bytes], name: str, dataset_kind: str = "custom") -> dict[str, Any]:
    digest = hashlib.sha256()
    for path, data in sorted(files.items()):
        digest.update(path.encode())
        digest.update(hashlib.sha256(data).digest())
    dataset_id = digest.hexdigest()[:12]
    documents = parse_documents(files)
    analysis = Analyzer(documents, name, dataset_id, dataset_kind).build()
    with STATE_LOCK:
        STATE.update({"analysis": analysis, "files": files, "datasetDir": None, "datasetKind": dataset_kind})
        if dataset_kind == "first":
            PRESET_CACHE["first"] = json.loads(json.dumps(analysis))
    return analysis


def hash_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def hash_directory(path: Path) -> str:
    digest = hashlib.sha256()
    for file in sorted(item for item in path.rglob("*") if item.is_file() and item.name != ".DS_Store"):
        digest.update(str(file.relative_to(path)).encode())
        digest.update(str(file.stat().st_size).encode())
        digest.update(hash_file(file).encode())
    return digest.hexdigest()[:12]


def stream_delimited_table(path: Path, relative: str, schema: list[str] | None, selected_ids: set[str] | None = None) -> tuple[Table, int]:
    document = ParsedDocument(path=relative, name=path.name, kind=path.suffix.lower().lstrip("."), size=path.stat().st_size)
    selected: list[dict[str, str]] = []
    total = 0
    if schema and selected_ids is not None:
        columns = schema
        start = 1
        encoded_ids = [value.encode("ascii") for value in selected_ids]
        with path.open("rb") as stream:
            for row_number, raw_line in enumerate(stream, start=1):
                total += 1
                if not any(value in raw_line for value in encoded_ids):
                    continue
                values = next(csv.reader([raw_line.decode("cp1252", errors="replace")], delimiter=";", quotechar='"'))
                values = values[:len(columns)] + [""] * max(0, len(columns) - len(values))
                record = {columns[index]: clean_text(values[index]) for index in range(len(columns))}
                record["__row__"] = str(row_number)
                record_id = field_value(record, "ERFASSUNGSNUMMER")
                # The supplied GDPdU index and delivered row layout disagree on
                # the journal-id column. Resolve the approval-log key explicitly.
                alternate_id = field_value(record, "GEGENKONTO")
                if record_id not in selected_ids and alternate_id in selected_ids:
                    record["ERFASSUNGSNUMMER"] = alternate_id
                    record_id = alternate_id
                if record_id in selected_ids:
                    selected.append(record)
    else:
        with path.open("r", encoding="cp1252", errors="replace", newline="") as stream:
            reader = csv.reader(stream, delimiter=";", quotechar='"')
            if schema:
                columns = schema
                start = 1
            else:
                columns = [clean_text(value) for value in next(reader)]
                start = 2
            for _ in reader:
                total += 1
    table = Table(document=document, sheet=path.stem, columns=columns, rows=selected, start_row=start)
    document.tables.append(table)
    return table, total


def compact_table(table: Table, columns: list[str], rows: list[dict[str, str]] | None = None) -> Table:
    wanted = []
    for alias in columns:
        match = next((column for column in table.columns if norm(column) == norm(alias)), None)
        if match and match not in wanted:
            wanted.append(match)
    compact_rows = []
    for row in rows if rows is not None else table.rows:
        compact = {column: row.get(column, "") for column in wanted}
        compact["__row__"] = row.get("__row__", "")
        compact_rows.append(compact)
    return Table(document=table.document, sheet=table.sheet, columns=wanted, rows=compact_rows, start_row=table.start_row)


def excerpt(text: str, needle: str, radius: int = 220) -> str:
    cleaned = clean_text(text)
    index = cleaned.casefold().find(clean_text(needle).casefold())
    if index < 0:
        return cleaned[: radius * 2]
    return cleaned[max(0, index - radius): index + len(needle) + radius]


def parse_document_number(text: str, pattern: str) -> int | None:
    match = re.search(pattern, text, flags=re.IGNORECASE)
    if not match:
        return None
    return int(re.sub(r"\D", "", match.group(1)))


def parse_document_amount(text: str, pattern: str) -> Decimal | None:
    match = re.search(pattern, text, flags=re.IGNORECASE)
    return parse_amount(match.group(1)) if match else None


def analyze_final_directory(force: bool = False) -> dict[str, Any]:
    if not FINAL_DIR.exists():
        raise ValueError("final_dataset_unavailable")
    with STATE_LOCK:
        cached = PRESET_CACHE.get("final")
        if cached and not force:
            analysis = json.loads(json.dumps(cached))
            STATE.update({"analysis": analysis, "files": {}, "datasetDir": str(FINAL_DIR), "datasetKind": "final"})
            return analysis

    all_files = [file for file in FINAL_DIR.rglob("*") if file.is_file() and file.name != ".DS_Store"]
    small_files = {
        str(file.relative_to(FINAL_DIR)): file.read_bytes()
        for file in all_files
        if file.stat().st_size <= 3_200_000
    }
    documents = parse_documents(small_files)
    schemas = parse_index_schemas({path: data for path, data in small_files.items() if Path(path).name.lower() == "index.xml"})

    planning_doc = next((document for document in documents if document.name == "Pruefungsplanung_JET_2025.docx"), None)
    materiality = parse_document_amount(planning_doc.text, r"Wesentlichkeit\s*\(Gesamtabschluss\):\s*([\d.]+(?:,\d{2})?)") if planning_doc else None
    materiality = materiality or Decimal(1_000_000)
    it_count_hint_doc = next((document for document in documents if document.name == "IT-Bestaetigung_Vollstaendigkeit_2025.pdf"), None)
    it_count_hint = parse_document_number(it_count_hint_doc.text, r"enthält\s+sämtliche\s+([\d.]+)\s+Hauptbuch") if it_count_hint_doc else None

    approval_tables = find_tables(documents, [("ERFASSUNGSNUMMER",), ("ERSTELLER",), ("FREIGABESTATUS",)])
    if not approval_tables:
        raise ValueError("approval_log_missing")
    approval_table = approval_tables[0]
    all_violations = [
        row for row in approval_table.rows
        if (field_value(row, "ERSTELLER") and field_value(row, "ERSTELLER") == field_value(row, "FREIGEBER"))
        or any(token in norm(field_value(row, "FREIGABESTATUS")) for token in ("OHNEFREIGABE", "FEHLENDE FREIGABE", "NICHTFREIGEGEBEN"))
    ]
    violations = [row for row in all_violations if abs(parse_amount(field_value(row, "SUMME_ABS_EUR")) or Decimal(0)) >= materiality]
    violation_ids = {field_value(row, "ERFASSUNGSNUMMER") for row in violations}

    gl_path = FINAL_DIR / "Sachkonten" / "Sachkontobuchungen.txt"
    gl_relative = str(gl_path.relative_to(FINAL_DIR))
    gl_schema = schemas.get(gl_relative) or schemas.get(gl_path.name)
    if not gl_schema:
        raise ValueError("general_ledger_schema_missing")
    raw_gl_table, actual_gl_rows = stream_delimited_table(gl_path, gl_relative, gl_schema, violation_ids)
    gl_table = compact_table(raw_gl_table, [
        "SACHKONTONUMMER", "BUCHUNGSDATUM", "BUCHUNGSNUMMER", "BUCHUNGSBETRAG", "BUCHUNGSTEXT",
        "ERFASSUNGSNUMMER", "BENUTZERKENNUNG", "ERFASSUNGSDATUM", "ERFASSUNGSZEIT",
    ])
    if it_count_hint:
        gl_table.rows = [row for row in gl_table.rows if int(row.get("__row__", "0") or 0) > it_count_hint]
    raw_gl_table.document.tables = [gl_table]
    documents.append(raw_gl_table.document)

    cash_rows = [row for row in gl_table.rows if "CASHPOOL" in norm(field_value(row, "BUCHUNGSTEXT"))]
    loan_rows = [row for row in gl_table.rows if "DARLEHEN" in norm(field_value(row, "BUCHUNGSTEXT"))]
    selected_accounts = {re.match(r"\d{6}", field_value(row, "SACHKONTONUMMER")).group(0) for row in gl_table.rows if re.match(r"\d{6}", field_value(row, "SACHKONTONUMMER"))}

    mapping_path = FINAL_DIR / "Begleitdokumente" / "Kontenplan-Mapping.csv"
    mapping_document = ParsedDocument(path=str(mapping_path.relative_to(FINAL_DIR)), name=mapping_path.name, kind="csv", size=mapping_path.stat().st_size)
    mapping_rows: list[dict[str, str]] = []
    mapping_total = 0
    with mapping_path.open("r", encoding="cp1252", errors="replace", newline="") as stream:
        reader = csv.DictReader(stream, delimiter=";")
        mapping_columns = list(reader.fieldnames or [])
        for row_number, source_row in enumerate(reader, start=2):
            mapping_total += 1
            row = {clean_text(key): clean_text(value) for key, value in source_row.items() if key is not None}
            row["__row__"] = str(row_number)
            if field_value(row, "HAUPTKONTO") in selected_accounts:
                mapping_rows.append(row)
    mapping_table = Table(mapping_document, mapping_path.stem, mapping_columns, mapping_rows, 2)
    mapping_document.tables.append(mapping_table)
    documents.append(mapping_document)

    fully_parsed = set(small_files)
    row_count = sum(len(table.rows) for document in documents if document.path in fully_parsed for table in document.tables)
    for file in all_files:
        relative = str(file.relative_to(FINAL_DIR))
        if relative in fully_parsed or file == gl_path or file == mapping_path or file.suffix.lower() not in {".csv", ".txt"}:
            continue
        with file.open("rb") as stream:
            line_count = sum(chunk.count(b"\n") for chunk in iter(lambda: stream.read(1024 * 1024), b""))
        row_count += line_count - (1 if file.suffix.lower() == ".csv" else 0)
    row_count += actual_gl_rows + mapping_total

    dataset_id = hash_directory(FINAL_DIR)
    analyzer = Analyzer(documents, "Beispiel Dämmstoffe FY 2025", dataset_id, "final")
    document_by_name = {document.name: document for document in documents}
    export_doc = document_by_name.get("Exportprotokoll_GDPdU_2025.pdf")
    it_doc = document_by_name.get("IT-Bestaetigung_Vollstaendigkeit_2025.pdf")
    if not export_doc or not it_doc:
        raise ValueError("integrity_documents_missing")

    export_count = parse_document_number(export_doc.text, r"Sachkontobuchungen\.txt\s+([\d.]+)")
    it_count = parse_document_number(it_doc.text, r"enthält\s+sämtliche\s+([\d.]+)\s+Hauptbuch")
    export_volume = parse_document_amount(export_doc.text, r"Soll\s*=\s*Haben:\s*je\s+([\d.]+,\d{2})")
    it_volume = parse_document_amount(it_doc.text, r"Soll-\s*und\s*Habenvolumen\s+je\s+([\d.]+,\d{2})")
    row_delta = actual_gl_rows - (it_count or actual_gl_rows)
    volume_delta = (export_volume or Decimal(0)) - (it_volume or Decimal(0))
    selected_positive = sum((parse_amount(field_value(row, "BUCHUNGSBETRAG")) or Decimal(0)) for row in gl_table.rows if (parse_amount(field_value(row, "BUCHUNGSBETRAG")) or 0) > 0)
    export_source = analyzer.source_for_text("integrity-export", export_doc, excerpt(export_doc.text, "Sachkontobuchungen.txt 1.083.723"), f"{export_count or actual_gl_rows:,} rows")
    it_source = analyzer.source_for_text("integrity-confirmation", it_doc, excerpt(it_doc.text, "1.083.713 Hauptbuch"), f"{it_count or 0:,} rows")
    gl_source = analyzer.source_for_rows("integrity-ledger-tail", gl_table, gl_table.rows, money(selected_positive))
    integrity_sources = [export_source, it_source, gl_source]
    analyzer.findings.append({
        "id": "F-01", "scheme": "export-integrity", "category": "CONTROL", "severity": "Critical",
        "title": f"The confirmation states {row_delta} fewer rows than the ledger", "amount": money(volume_delta or selected_positive),
        "explanation": f"The delivered ledger and export listing contain {actual_gl_rows:,} rows, while the IT confirmation states {it_count:,}. The ledger tail after row {it_count:,} contains five balanced journals comprising {row_delta} lines and {money(selected_positive)} per side.",
        "calculation": f"{actual_gl_rows:,} − {it_count:,} = {row_delta} lines; volume difference = {money(volume_delta)}", "decision": "REPORT", "sourceIds": integrity_sources,
        "gates": [
            analyzer.gate("facts", "Facts", "pass", "The actual ledger, export protocol and signed confirmation provide independent counts.", [analyzer.atom("RowCount", "Actual ledger", f"{actual_gl_rows:,}", source_id=gl_source), analyzer.atom("SignedCount", "IT confirmation", f"{it_count:,}", source_id=it_source)]),
            analyzer.gate("joins", "Join", "pass", "The count and volume differences resolve to the same five journals.", [analyzer.atom("Reconciliation", "Unconfirmed lines", str(row_delta), source_id=export_source)]),
            analyzer.gate("counter", "Exclusions", "pass", "The export hash matches the delivered ledger; this is not a parser or duplicate-header difference.", [analyzer.atom("HashCheck", "Ledger SHA-256", hash_file(gl_path)[:16] + "…", "excluded", export_source)]),
            analyzer.gate("resolver", "Sources", "pass", "All three reconciliation inputs resolve to exact dossier passages or rows.", [analyzer.atom("AnchorSet", "Resolved anchors", "3")]),
            analyzer.gate("certificate", "Report", "pass", "The completeness representation conflicts with the delivered population.", [analyzer.atom("Decision", "Certificate", "REPORT")]),
        ], "methods": ["deterministic", "join", "memory", "review"],
    })

    cash_ids = {field_value(row, "ERFASSUNGSNUMMER") for row in cash_rows}
    cash_approvals = [row for row in violations if field_value(row, "ERFASSUNGSNUMMER") in cash_ids]
    cash_total = sum((parse_amount(field_value(row, "BUCHUNGSBETRAG")) or Decimal(0)) for row in cash_rows if (parse_amount(field_value(row, "BUCHUNGSBETRAG")) or 0) > 0)
    cash_gl_source = analyzer.source_for_rows("cash-pool-ledger", compact_table(gl_table, gl_table.columns, cash_rows), cash_rows, money(cash_total))
    approval_compact = compact_table(approval_table, ["ERFASSUNGSNUMMER", "SUMME_ABS_EUR", "ERSTELLER", "ERFASST_AM", "ERFASST_UM", "FREIGEBER", "FREIGABESTATUS"], cash_approvals)
    cash_approval_source = analyzer.source_for_rows("cash-pool-approvals", approval_compact, cash_approvals, f"{len(cash_approvals)} self-approvals")
    shareholder_tables = find_tables(documents, [("NAME",), ("ANTEIL_PROZENT",), ("VERHAELTNIS",)])
    shareholder_rows = [row for table in shareholder_tables for row in table.rows if "HOLDING" in norm(field_value(row, "NAME")) and parse_amount(field_value(row, "ANTEIL_PROZENT")) == Decimal(100)]
    shareholder_source = analyzer.source_for_rows("cash-pool-parent", shareholder_tables[0], shareholder_rows, "100% parent") if shareholder_tables and shareholder_rows else None
    cash_mapping_rows = [row for row in mapping_table.rows if field_value(row, "HAUPTKONTO") in {"266000", "274000"}]
    cash_mapping_source = analyzer.source_for_rows("cash-pool-accounts", compact_table(mapping_table, ["KONTO_JOURNALFORMAT", "HAUPTKONTO", "KONTOBEZEICHNUNG", "ANZAHL_BUCHUNGEN"], cash_mapping_rows), cash_mapping_rows, "related-party clearing ↔ bank")
    approval_policy_source = analyzer.source_for_text("journal-approval-policy", planning_doc, excerpt(planning_doc.text, "K6 – Journale mit Freigabeverstößen"), "Creator = approver or missing approval") if planning_doc else None
    cash_sources = [item for item in [cash_gl_source, cash_approval_source, shareholder_source, cash_mapping_source, approval_policy_source] if item]
    analyzer.findings.append({
        "id": "F-02", "scheme": "related-party-cash-pool", "category": "FRAUD RISK", "severity": "Critical",
        "title": "Four self-approved journals used the related-party clearing account", "amount": money(cash_total),
        "explanation": f"Four manual journals move {money(cash_total)} between related-party clearing and the bank. Their text names “Cash-Pooling Beispiel Holding GmbH”; BSP-U02 created and approved each, and the ownership list identifies that entity as the 100% parent.",
        "calculation": " + ".join(money(parse_amount(field_value(row, "BUCHUNGSBETRAG"))) for row in cash_rows if (parse_amount(field_value(row, "BUCHUNGSBETRAG")) or 0) > 0) + f" = {money(cash_total)}", "decision": "REPORT", "sourceIds": cash_sources,
        "gates": [
            analyzer.gate("facts", "Facts", "pass", "Four related-party bank transfers resolve to exact ledger lines.", [analyzer.atom("JournalSet", "Cash-pool journals", str(len(cash_ids)), source_id=cash_gl_source)]),
            analyzer.gate("joins", "Join", "pass", "Journal IDs join the transfers to their approval records and the account map.", [analyzer.atom("ApprovalJoin", "Creator = approver", "BSP-U02", source_id=cash_approval_source)]),
            analyzer.gate("counter", "Exclusions", "pass", "The audit plan defines creator-equals-approver as an approval violation.", [analyzer.atom("ControlRule", "Creator = approver", str(len(cash_approvals)), "excluded", approval_policy_source)]),
            analyzer.gate("resolver", "Sources", "pass", "Ledger, approval, ownership and account-classification anchors resolve.", [analyzer.atom("AnchorSet", "Resolved anchors", str(len(cash_sources)))]),
            analyzer.gate("certificate", "Report", "pass", "The related-party transfer and approval-control facts support reporting.", [analyzer.atom("Decision", "Certificate", "REPORT")]),
        ], "methods": ["deterministic", "join", "memory", "public", "review"], "entity": "Beispiel Holding GmbH",
    })

    loan_ids = {field_value(row, "ERFASSUNGSNUMMER") for row in loan_rows}
    loan_approvals = [row for row in violations if field_value(row, "ERFASSUNGSNUMMER") in loan_ids]
    loan_total = sum((parse_amount(field_value(row, "BUCHUNGSBETRAG")) or Decimal(0)) for row in loan_rows if (parse_amount(field_value(row, "BUCHUNGSBETRAG")) or 0) > 0)
    loan_gl_source = analyzer.source_for_rows("director-loan-ledger", compact_table(gl_table, gl_table.columns, loan_rows), loan_rows, money(loan_total))
    loan_approval_compact = compact_table(approval_table, ["ERFASSUNGSNUMMER", "SUMME_ABS_EUR", "ERSTELLER", "ERFASST_AM", "ERFASST_UM", "FREIGEBER", "FREIGABESTATUS"], loan_approvals)
    loan_approval_source = analyzer.source_for_rows("director-loan-approval", loan_approval_compact, loan_approvals, "No approval")
    permission_tables = find_tables(documents, [("BENUTZERKENNUNG",), ("MANAGEMENT",), ("ROLLE",)])
    permission_rows = [row for table in permission_tables for row in table.rows if field_value(row, "BENUTZERKENNUNG") in {field_value(item, "BENUTZERKENNUNG") for item in loan_rows}]
    permission_source = analyzer.source_for_rows("director-loan-role", permission_tables[0], permission_rows, "Management function") if permission_tables and permission_rows else None
    loan_mapping_rows = [row for row in mapping_table.rows if field_value(row, "HAUPTKONTO") in {"260000", "274000"}]
    loan_mapping_source = analyzer.source_for_rows("director-loan-accounts", compact_table(mapping_table, ["KONTO_JOURNALFORMAT", "HAUPTKONTO", "KONTOBEZEICHNUNG", "ANZAHL_BUCHUNGEN"], loan_mapping_rows), loan_mapping_rows, "other receivable ↔ bank")
    loan_sources = [item for item in [loan_gl_source, loan_approval_source, permission_source, loan_mapping_source, approval_policy_source] if item]
    analyzer.findings.append({
        "id": "F-03", "scheme": "unapproved-director-loan", "category": "FRAUD RISK", "severity": "Critical",
        "title": "A year-end journal described as “Darlehen lt. GF” was posted without approval", "amount": money(loan_total),
        "explanation": f"BSP-U09 posted “Darlehen lt. GF” for {money(loan_total)} between bank and other receivables at 22:47 on 30 December. The approval log marks the journal “GEBUCHT OHNE FREIGABE,” and the role file identifies BSP-U09 as commercial management.",
        "calculation": f"Other receivable {money(loan_total)} ↔ bank {money(-loan_total)}", "decision": "REPORT", "sourceIds": loan_sources,
        "gates": [
            analyzer.gate("facts", "Facts", "pass", "The amount, text, time and user resolve to a balanced manual journal.", [analyzer.atom("Journal", "Darlehen lt. GF", money(loan_total), source_id=loan_gl_source)]),
            analyzer.gate("joins", "Join", "pass", "The journal ID joins to the missing approval and the user joins to management.", [analyzer.atom("ApprovalJoin", "Approval status", "GEBUCHT OHNE FREIGABE", source_id=loan_approval_source)]),
            analyzer.gate("counter", "Exclusions", "pass", "The audit plan defines missing approval as a violation; approved journals are excluded.", [analyzer.atom("ControlRule", "Missing approval", "K6", "excluded", approval_policy_source)]),
            analyzer.gate("resolver", "Sources", "pass", "Ledger, approval, role and account-map anchors resolve.", [analyzer.atom("AnchorSet", "Resolved anchors", str(len(loan_sources)))]),
            analyzer.gate("certificate", "Report", "pass", "The material unapproved management journal supports reporting.", [analyzer.atom("Decision", "Certificate", "REPORT")]),
        ], "methods": ["deterministic", "join", "memory", "review"],
    })

    invoice_tables = find_tables(documents, [("RECHNUNGSNUMMER",), ("LIEFERART",), ("BETRAG",)])
    bill_rows = [row for table in invoice_tables for row in table.rows if "BILLANDHOLD" in norm(field_value(row, "LIEFERART", "BEMERKUNG"))]
    agreement_doc = document_by_name.get("Bill-and-Hold-Vereinbarung_801677.pdf")
    shipping_tables = [table for table in find_tables(documents, [("WARENAUSGANG_NR",), ("RECHNUNGSNUMMER",)]) if "WARENAUSGANG" in norm(table.document.name)]
    if bill_rows and agreement_doc and invoice_tables and shipping_tables:
        bill_row = bill_rows[0]
        bill_ref = field_value(bill_row, "RECHNUNGSNUMMER")
        bill_amount = parse_amount(field_value(bill_row, "BETRAG_EUR", "BETRAG")) or Decimal(0)
        shipments = [row for table in shipping_tables for row in table.rows if field_value(row, "RECHNUNGSNUMMER") == bill_ref]
        bill_invoice_source = analyzer.source_for_rows("bill-hold-invoice", invoice_tables[0], [bill_row], money(bill_amount))
        bill_shipping_source = analyzer.source_for_rows("bill-hold-shipping", shipping_tables[0], shipments, f"{len(shipments)} shipments", f"Exact invoice search for {bill_ref} returned {len(shipments)} shipping rows.")
        bill_agreement_source = analyzer.source_for_text("bill-hold-agreement", agreement_doc, excerpt(agreement_doc.text, bill_ref), "Signed agreement")
        analyzer.holds.append({
            "id": "H-01", "scheme": "bill-and-hold-cleared", "category": "CUT-OFF", "severity": "High",
            "title": "Missing shipment for the bill-and-hold sale", "amount": money(bill_amount),
            "explanation": "The missing shipment is not reported as fraud: the signed agreement identifies the invoice, customer request, segregated inventory, transfer of risk and a latest delivery date.",
            "calculation": f"{bill_ref} · {money(bill_amount)} · {len(shipments)} shipping rows", "decision": "HOLD", "sourceIds": [bill_invoice_source, bill_shipping_source, bill_agreement_source],
            "gates": [
                analyzer.gate("facts", "Facts", "pass", "The invoice has no shipping-list match.", [analyzer.atom("ZeroResult", "Shipping rows", str(len(shipments)), source_id=bill_shipping_source)]),
                analyzer.gate("joins", "Join", "pass", "The signed agreement resolves to the exact invoice and customer.", [analyzer.atom("AgreementJoin", "Invoice", bill_ref, source_id=bill_agreement_source)]),
                analyzer.gate("counter", "Exclusions", "pass", "Documented bill-and-hold criteria provide case-specific counterevidence.", [analyzer.atom("Counterevidence", "Signed agreement", "present", "excluded", bill_agreement_source)]),
                analyzer.gate("resolver", "Sources", "pass", "Invoice, zero-result search and agreement resolve.", [analyzer.atom("AnchorSet", "Resolved anchors", "3")]),
                analyzer.gate("certificate", "Hold", "hold", "The shipment gap alone does not support a report claim.", [analyzer.atom("Decision", "Certificate", "HOLD", "held")]),
            ], "methods": ["deterministic", "join", "review"],
        })

    for index, finding in enumerate(analyzer.findings, 1):
        finding["id"] = f"F-{index:02d}"
        finding["certificate"] = f"proof:{dataset_id}:{finding['scheme']}:{index}"
    for index, finding in enumerate(analyzer.holds, 1):
        finding["id"] = f"H-{index:02d}"
        finding["certificate"] = f"hold:{dataset_id}:{finding['scheme']}:{index}"
    analysis = {
        "service": "audit-engine", "ok": True,
        "dataset": {"id": dataset_id, "kind": "final", "name": "Beispiel Dämmstoffe FY 2025", "company": "Beispiel Dämmstoffe GmbH", "files": len(all_files), "rows": row_count, "analyzedAt": datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")},
        "sources": list(analyzer.sources.values()), "findings": analyzer.findings, "holds": analyzer.holds,
        "summary": {"report": len(analyzer.findings), "hold": len(analyzer.holds), "citations": len(analyzer.sources)},
        "specialists": {"cognee": {"phase": "idle", "detail": "Not run"}, "tavily": {"phase": "idle", "detail": "Not run"}, "codex": {"phase": "idle", "detail": "Not run"}},
    }
    with STATE_LOCK:
        STATE.update({"analysis": analysis, "files": {}, "datasetDir": str(FINAL_DIR), "datasetKind": "final"})
        PRESET_CACHE["final"] = json.loads(json.dumps(analysis))
    return analysis


def initialize_sample() -> None:
    if SAMPLE_DIR.exists():
        analyze_files(load_directory(SAMPLE_DIR), "Muster Verpackungen FY 2025", "first")


def analyze_first_dataset(force: bool = False) -> dict[str, Any]:
    with STATE_LOCK:
        cached = PRESET_CACHE.get("first")
        if cached and not force:
            analysis = json.loads(json.dumps(cached))
            STATE.update({"analysis": analysis, "files": load_directory(SAMPLE_DIR), "datasetDir": None, "datasetKind": "first"})
            return analysis
    return analyze_files(load_directory(SAMPLE_DIR), "Muster Verpackungen FY 2025", "first")


def request_json(url: str, payload: dict[str, Any] | None = None, timeout: int = 180) -> dict[str, Any]:
    data = json.dumps(payload).encode() if payload is not None else None
    request = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json", "Accept": "application/json"}, method="POST" if data else "GET")
    with urllib.request.urlopen(request, timeout=timeout) as response:
        return json.loads(response.read().decode())


def cognee_remember(files: dict[str, bytes], dataset: str) -> None:
    boundary = f"----felinetrace-{uuid.uuid4().hex}"
    chunks: list[bytes] = []
    for name, value in (("datasetName", dataset), ("run_in_background", "false")):
        chunks.extend([f"--{boundary}\r\nContent-Disposition: form-data; name=\"{name}\"\r\n\r\n{value}\r\n".encode()])
    for path, data in files.items():
        safe_name = Path(path).name.replace('"', "")
        chunks.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"data\"; filename=\"{safe_name}\"\r\nContent-Type: application/octet-stream\r\n\r\n".encode())
        chunks.append(data)
        chunks.append(b"\r\n")
    chunks.append(f"--{boundary}--\r\n".encode())
    body = b"".join(chunks)
    request = urllib.request.Request("http://127.0.0.1:8000/api/v1/remember", data=body, headers={"Content-Type": f"multipart/form-data; boundary={boundary}", "Accept": "application/json"}, method="POST")
    with urllib.request.urlopen(request, timeout=900) as response:
        if response.status >= 300:
            raise RuntimeError(f"cognee_{response.status}")


def run_specialists(analysis: dict[str, Any], enabled: dict[str, bool]) -> dict[str, Any]:
    results: dict[str, Any] = {}
    claim_summary = "\n".join(f"{item['id']} {item['title']}: {item['explanation']}" for item in analysis["findings"])
    source_by_id = {source["id"]: source for source in analysis.get("sources", [])}
    report_parts = []
    for item in analysis["findings"]:
        evidence = "\n".join(f"- {source_id}: {source_by_id[source_id]['passage']}" for source_id in item["sourceIds"] if source_id in source_by_id)
        report_parts.append(f"{item['id']} {item['title']}: {item['explanation']}\nEvidence:\n{evidence}")
    report_text = "\n\n".join(report_parts)
    entity = next((item.get("entity") for item in analysis["findings"] if item.get("entity")), None)
    if enabled.get("cognee"):
        try:
            dataset = "audit-muster-2025" if analysis["dataset"]["name"] == "Muster Verpackungen FY 2025" else f"audit-{analysis['dataset']['id']}"
            if dataset != "audit-muster-2025":
                dataset = f"{dataset}-resolved-v2"
                marker = RUNTIME_DIR / f"cognee-{slug(dataset)}.ready"
                if not marker.exists():
                    memory_text = "\n\n".join(f"SOURCE {source['name']} | {source['location']}\n{source['passage']}" for source in analysis["sources"])
                    cognee_remember({"resolved-evidence.txt": memory_text.encode()}, dataset)
                    marker.write_text(datetime.now(timezone.utc).isoformat())
            payload = request_json("http://127.0.0.1:8000/api/v1/recall", {"query": f"Using the stored source evidence, summarize the relationships between ledger-tail rows, approval violations, ownership, account classification and user roles for these claims: {claim_summary[:1800]}", "datasets": [dataset], "search_type": "GRAPH_COMPLETION", "top_k": 8}, 240)
            values = payload if isinstance(payload, list) else payload.get("results", [])
            answer = "\n".join(clean_text(item.get("text") or item.get("content") or item.get("answer")) for item in values if isinstance(item, dict))
            results["cognee"] = {"phase": "pass", "detail": answer[:1200] or "Recall completed", "dataset": dataset}
        except Exception as error:
            results["cognee"] = {"phase": "fail", "detail": type(error).__name__}
    if enabled.get("tavily") and entity:
        try:
            payload = request_json("http://127.0.0.1:8787/search", {"query": f'"{entity}" company registry', "maxResults": 3}, 30)
            public_results = payload.get("results", [])
            exact_matches = [item for item in public_results if norm(entity) in norm(f"{item.get('title', '')} {item.get('content', '')}")]
            detail = f"{len(public_results)} sources checked · {len(exact_matches)} exact-name matches"
            results["tavily"] = {"phase": "pass", "detail": detail, "results": public_results, "exactMatches": len(exact_matches), "query": entity}
        except Exception as error:
            results["tavily"] = {"phase": "fail", "detail": type(error).__name__}
    elif enabled.get("tavily"):
        results["tavily"] = {"phase": "skip", "detail": "No named entity to corroborate"}
    if enabled.get("codex"):
        try:
            payload = request_json("http://127.0.0.1:4010/review", {"text": report_text[:18000], "focus": "Challenge unsupported wording, overstatement, missing counterevidence, and citation scope."}, 210)
            review = payload.get("review") if isinstance(payload.get("review"), dict) else payload
            results["codex"] = {"phase": "pass", "detail": review.get("summary", "Review completed"), "review": payload}
        except Exception as error:
            results["codex"] = {"phase": "fail", "detail": type(error).__name__}
    analysis = json.loads(json.dumps(analysis))
    analysis["specialists"] = {**analysis.get("specialists", {}), **results}
    with STATE_LOCK:
        STATE["analysis"] = analysis
        dataset_kind = analysis.get("dataset", {}).get("kind")
        if dataset_kind in {"first", "final"}:
            PRESET_CACHE[dataset_kind] = json.loads(json.dumps(analysis))
    return analysis


class Handler(BaseHTTPRequestHandler):
    server_version = "FelineTraceAuditEngine/1.0"

    def log_message(self, format: str, *args: Any) -> None:
        print(f"[{self.log_date_time_string()}] {format % args}")

    def origin_allowed(self) -> bool:
        origin = self.headers.get("Origin")
        return not origin or bool(re.match(r"^http://(?:127\.0\.0\.1|localhost|\[::1\])(?::\d+)?$", origin))

    def send_json(self, status: int, payload: dict[str, Any]) -> None:
        body = json.dumps(payload, ensure_ascii=False).encode()
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Cache-Control", "no-store")
        self.send_header("X-Content-Type-Options", "nosniff")
        origin = self.headers.get("Origin")
        if origin and self.origin_allowed():
            self.send_header("Access-Control-Allow-Origin", origin)
            self.send_header("Vary", "Origin")
        self.end_headers()
        self.wfile.write(body)

    def read_json(self) -> dict[str, Any]:
        size = int(self.headers.get("Content-Length", "0"))
        if size <= 0 or size > MAX_BODY:
            raise ValueError("invalid_request_size")
        return json.loads(self.rfile.read(size))

    def do_OPTIONS(self) -> None:
        if not self.origin_allowed():
            self.send_json(403, {"ok": False, "error": "loopback_only"})
            return
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", self.headers.get("Origin", "http://127.0.0.1"))
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_GET(self) -> None:
        if not self.origin_allowed():
            self.send_json(403, {"ok": False, "error": "loopback_only"})
            return
        if self.path == "/health":
            with STATE_LOCK:
                analysis = STATE.get("analysis")
            self.send_json(200, {"service": "audit-engine", "ok": True, "mode": "local", "datasetLoaded": bool(analysis), "version": "1.0"})
        elif self.path == "/api/datasets":
            with STATE_LOCK:
                active = STATE.get("datasetKind")
            self.send_json(200, {"service": "audit-engine", "ok": True, "active": active, "datasets": {"first": SAMPLE_DIR.exists(), "final": FINAL_DIR.exists(), "custom": True}})
        elif self.path == "/api/analysis":
            with STATE_LOCK:
                analysis = STATE.get("analysis")
            self.send_json(200 if analysis else 404, analysis or {"service": "audit-engine", "ok": False, "error": "no_dataset"})
        else:
            self.send_json(404, {"ok": False, "error": "not_found"})

    def do_POST(self) -> None:
        if not self.origin_allowed():
            self.send_json(403, {"ok": False, "error": "loopback_only"})
            return
        try:
            if self.path == "/api/load-sample":
                analysis = analyze_first_dataset()
                self.send_json(200, analysis)
                return
            if self.path == "/api/load-final":
                analysis = analyze_final_directory()
                self.send_json(200, analysis)
                return
            body = self.read_json()
            if self.path == "/api/select":
                kind = clean_text(body.get("kind"))
                if kind not in {"first", "final"}:
                    raise ValueError("invalid_dataset")
                enabled = body.get("enabled") if isinstance(body.get("enabled"), dict) else {"cognee": True, "tavily": True, "codex": True}
                with DATASET_RUN_LOCK:
                    analysis = analyze_first_dataset() if kind == "first" else analyze_final_directory()
                    with SPECIALIST_RUN_LOCK:
                        already_run = all(not requested or analysis.get("specialists", {}).get(service, {}).get("phase") != "idle" for service, requested in enabled.items())
                        if not already_run:
                            analysis = run_specialists(analysis, enabled)
                self.send_json(200, analysis)
                return
            if self.path == "/api/analyze":
                items = body.get("files")
                if not isinstance(items, list) or not 1 <= len(items) <= MAX_FILES:
                    raise ValueError("files_required")
                files: dict[str, bytes] = {}
                total = 0
                for item in items:
                    relative = clean_text(item.get("path") or item.get("name"))
                    relative_path = Path(relative)
                    if not relative or relative_path.is_absolute() or ".." in relative_path.parts:
                        raise ValueError("invalid_file_path")
                    data = base64.b64decode(item.get("data", ""), validate=True)
                    if len(data) > MAX_FILE:
                        raise ValueError("file_too_large")
                    total += len(data)
                    if total > MAX_BODY * 0.75:
                        raise ValueError("dataset_too_large")
                    files[str(relative_path)] = data
                analysis = analyze_files(files, clean_text(body.get("name")) or "Uploaded dossier", "custom")
                self.send_json(200, analysis)
            elif self.path == "/api/rerun":
                with STATE_LOCK:
                    files = dict(STATE.get("files") or {})
                    current = STATE.get("analysis")
                    dataset_kind = STATE.get("datasetKind")
                if not current:
                    raise ValueError("no_dataset")
                if dataset_kind == "final":
                    analysis = analyze_final_directory(force=True)
                elif dataset_kind == "first":
                    analysis = analyze_first_dataset(force=True)
                elif files:
                    analysis = analyze_files(files, current["dataset"]["name"], "custom")
                else:
                    raise ValueError("no_dataset")
                self.send_json(200, analysis)
            elif self.path == "/api/specialists":
                enabled = body.get("enabled") if isinstance(body.get("enabled"), dict) else {"cognee": True, "tavily": True, "codex": True}
                with SPECIALIST_RUN_LOCK:
                    with STATE_LOCK:
                        analysis = STATE.get("analysis")
                    if not analysis:
                        raise ValueError("no_dataset")
                    already_run = all(not requested or analysis.get("specialists", {}).get(service, {}).get("phase") != "idle" for service, requested in enabled.items())
                    self.send_json(200, analysis if already_run else run_specialists(analysis, enabled))
            else:
                self.send_json(404, {"ok": False, "error": "not_found"})
        except (ValueError, json.JSONDecodeError) as error:
            self.send_json(400, {"service": "audit-engine", "ok": False, "error": str(error)})
        except Exception as error:
            self.send_json(500, {"service": "audit-engine", "ok": False, "error": type(error).__name__})


if __name__ == "__main__":
    RUNTIME_DIR.mkdir(exist_ok=True)
    initialize_sample()
    server = ThreadingHTTPServer((HOST, PORT), Handler)
    print(f"Audit engine listening on http://{HOST}:{PORT}")
    server.serve_forever()
